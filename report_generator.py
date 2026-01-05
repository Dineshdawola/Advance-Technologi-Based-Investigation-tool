from docx import Document # Legal python alternative to DocX.dll
import datetime

def generate_forensic_report(case_data, suspect_number):
    doc = Document()
    doc.add_heading(f'Forensic Analysis Report: {suspect_number}', 0)
    
    # Adding Metadata
    doc.add_paragraph(f"Report Date: {datetime.datetime.now()}")
    doc.add_paragraph(f"Source Database: dlt.db / CDR.db")
    
    # Summary Table logic
    table = doc.add_table(rows=1, cols=3)
    # logic to populate forensic findings...
    
    doc.save(f"Reports/Report_{suspect_number}.docx")